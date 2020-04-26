from datetime import datetime , date, timedelta
import datetime as dt
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.db.models import Count, F
from django.http import JsonResponse, HttpResponse, HttpResponseNotFound

from DMS.utils import WithChoices
from work.models import Work
from .models import Document, InternalDoc, ExternalDoc
from .forms import ExternalDocForm, ExternalDocFilterForm, InternalDocFilterForm
from django.shortcuts import redirect, render

from docx import Document as docx
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches


@login_required(login_url='login')
def index(request):
    if request.user.username == 'readonly':
        return redirect('raw')
    works = Work.objects.all().order_by('-id')[:10]
    work_cnt = Work.objects.all().count()
    work_cnt_cr = Work.objects.filter(type='CR').count()
    work_cnt_e = Work.objects.filter(type='E').count()
    work_cnt_ca = Work.objects.filter(type='CA').count()

    work_cnt_cr_w = Work.objects.filter(create_date__gte=datetime.now() - timedelta(days=7)).count()
    work_cnt_com_w = Work.objects.exclude(complete_date__isnull=True).filter(create_date__gte=datetime.now() - timedelta(days=7)).count()
    work_cnt_cr_d = Work.objects.filter(create_date__range=(datetime.combine(datetime.now(), dt.time.min),
                                                                       datetime.combine(datetime.now(), dt.time.max))).count()
    work_cnt_com_d = Work.objects.exclude(complete_date__isnull=True).filter(create_date__range=(datetime.combine(datetime.now(), dt.time.min),
                                                                       datetime.combine(datetime.now(), dt.time.max))).count()

    internal = InternalDoc.objects.all().order_by('-id')[:10]
    internal_dashboard = InternalDoc.objects.filter(state='IN').order_by('-id')[:10]
    internal_dashboard_re = InternalDoc.objects.filter(state='RE').order_by('-id')[:10]
    internal_cnt = InternalDoc.objects.all().count()
    internal_cnt_in = InternalDoc.objects.filter(state='IN').count()
    internal_cnt_re = InternalDoc.objects.filter(state='RE').count()
    internal_cnt_ob = InternalDoc.objects.filter(state='OB').count()

    context = {
        'works': works,
        'work_cnt': work_cnt,
        'work_cnt_cr': work_cnt_cr,
        'work_cnt_e': work_cnt_e,
        'work_cnt_ca': work_cnt_ca,
        'work_cnt_cr_w': work_cnt_cr_w,
        'work_cnt_com_w': work_cnt_com_w,
        'work_cnt_cr_d': work_cnt_cr_d,
        'work_cnt_com_d': work_cnt_com_d,

        'documents': internal,
        'documents_in': internal_dashboard,
        'documents_re': internal_dashboard_re,
        'doc_cnt': internal_cnt,
        'doc_cnt_in': internal_cnt_in,
        'doc_cnt_re': internal_cnt_re,
        'doc_cnt_ob': internal_cnt_ob
    }
    return render(request, 'index.html', context=context)


@login_required(login_url='login')
def document_list(request, doc_type):
    if doc_type == 'external':
        filter_forms = ExternalDocFilterForm(request.GET)
        if filter_forms.is_valid():
            documents = ExternalDoc.objects.filter(
                name__icontains=filter_forms.cleaned_data['name'],
                source__icontains=filter_forms.cleaned_data['source'],
            )

            if filter_forms.cleaned_data['creator'] is not None:
                documents = documents.filter(creator=filter_forms.cleaned_data['creator'])

    elif doc_type == 'internal':
        filter_forms = InternalDocFilterForm(request.GET)
        if filter_forms.is_valid():
            documents = InternalDoc.objects.filter(
                name__icontains=filter_forms.cleaned_data['name'],
            ).prefetch_related('creator__department', 'creator__user', 'parent_doc')

            if filter_forms.cleaned_data['parent_doc'] is not None:
                documents = documents.filter(parent_doc=filter_forms.cleaned_data['parent_doc'])

            if filter_forms.cleaned_data['version'] is not None:
                documents = documents.filter(version=filter_forms.cleaned_data['version'])

            if filter_forms.cleaned_data['running_no'] is not None:
                documents = documents.filter(running_no=filter_forms.cleaned_data['running_no'])

            if filter_forms.cleaned_data['type'] != '':
                documents = documents.filter(type__exact=filter_forms.cleaned_data['type'])

            if filter_forms.cleaned_data['state'] != '':
                documents = documents.filter(state__exact=filter_forms.cleaned_data['state'])

            if filter_forms.cleaned_data['creator'] is not None:
                documents = documents.filter(creator=filter_forms.cleaned_data['creator'])

            if filter_forms.cleaned_data['department'] is not None:
                documents = documents.filter(creator__department=filter_forms.cleaned_data['department'])
        else:
            documents = InternalDoc.objects.all().prefetch_related('creator__department', 'creator__user', 'parent_doc')
    documents = documents.annotate(dept_name=F('creator__department__name'))

    page = request.GET.get('page', 1)
    paginator = Paginator(documents, 10)
    try:
        paginated_documents = paginator.page(page)
    except PageNotAnInteger:
        paginated_documents = paginator.page(1)
    except EmptyPage:
        paginated_documents = paginator.page(paginator.num_pages)

    context = {
        'documents': paginated_documents,
        'doc_type': doc_type,
        'filter_forms': filter_forms
    }
    return render(request, 'document_list.html', context=context)


@login_required(login_url='login')
def document_detail(request, id):
    document = Document.objects.get(pk=id)
    if hasattr(document, 'internaldoc'):
        return render(request, 'document_detail.html', {'document': document.internaldoc, 'doc_type': 'internal'})
    elif hasattr(document, 'externaldoc'):
        return render(request, 'document_detail.html', {'document': document.externaldoc, 'doc_type': 'external'})


@login_required(login_url='login')
def external_add(request):
    if request.method == 'POST':
        form = ExternalDocForm(request.POST, request.FILES)
        if form.is_valid():
            doc = form.save(commit=False)
            doc.creator = request.user.employee
            doc.save()
            return redirect('doc_detail', id=doc.id)
    return render(request, 'external_add.html', {'form': ExternalDocForm()})


def parse_html_time(time_string):
    return datetime.strptime(time_string, '%Y-%m-%dT%H:%M')


@login_required
def get_dashboard_work_cnt(request):
    work_cnt = Work.objects.all().count()
    work_cnt_cr = Work.objects.filter(type='CR').count()
    work_cnt_e = Work.objects.filter(type='E').count()
    work_cnt_ca = Work.objects.filter(type='CA').count()
    work_cnt_cr_w = Work.objects.filter(create_date__gte=datetime.now() - timedelta(days=7)).count()
    work_cnt_com_w = Work.objects.exclude(complete_date__isnull=True).filter(create_date__gte=datetime.now() - timedelta(days=7)).count()
    work_cnt_cr_d = Work.objects.filter(create_date__range=(datetime.combine(datetime.now(), dt.time.min),
                                                                       datetime.combine(datetime.now(), dt.time.max))).count()
    work_cnt_com_d = Work.objects.exclude(complete_date__isnull=True).filter(create_date__range=(datetime.combine(datetime.now(), dt.time.min),
                                                                        datetime.combine(datetime.now(), dt.time.max))).count()

    work_cnt = {
        "cnt": work_cnt,
        "cr": work_cnt_cr,
        "e": work_cnt_e,
        "ca": work_cnt_ca,
        'cr_w': work_cnt_cr_w,
        'com_w': work_cnt_com_w,
        'cr_d': work_cnt_cr_d,
        'com_d': work_cnt_com_d
    }
    return JsonResponse(work_cnt, safe=False)


@login_required
def get_dashboard_internal_cnt(request):
    internal_cnt = InternalDoc.objects.all().count()
    internal_cnt_in = InternalDoc.objects.filter(state='IN').count()
    internal_cnt_re = InternalDoc.objects.filter(state='RE').count()
    internal_cnt_ob = InternalDoc.objects.filter(state='OB').count()
    internal_cnt = {
        "cnt": internal_cnt,
        "in": internal_cnt_in,
        "re": internal_cnt_re,
        "ob": internal_cnt_ob
    }
    return JsonResponse(internal_cnt, safe=False)


@login_required
def get_dashboard_work_list(request):
    works = Work.objects.all().order_by('-id')[:10].annotate(
        document_name=F('document__name'),
        type_name=WithChoices(Work, 'type'),
        state_name=WithChoices(Work, 'state'),
    ).values()
    work_list = list(works)
    return JsonResponse(work_list, safe=False)


@login_required
def get_dashboard_internal_list(request):
    internal = InternalDoc.objects.all().order_by('-id')[:10].annotate(
        type_name=WithChoices(InternalDoc, 'type'),
        state_name=WithChoices(InternalDoc, 'state')
    ).values()
    internal_list = list(internal)
    return JsonResponse(internal_list, safe=False)


@login_required
def get_dashboard_internal_progress(request):
    progress = InternalDoc.objects.filter(
        state='IN'
    ).order_by('-id')[:10].values()
    progress_list = list(progress)
    return JsonResponse(progress_list, safe=False)


@login_required
def get_dashboard_internal_released(request):
    progress = InternalDoc.objects.filter(
        state='RE'
    ).order_by('-id')[:10].values()
    progress_list = list(progress)
    return JsonResponse(progress_list, safe=False)


@login_required(login_url='login')
def document_template(request, id):

    doc = InternalDoc.objects.get(pk=id)

    document = docx()

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('LogoStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(60)
    obj_font.name = 'Arial'

    p = document.add_paragraph()
    p.add_run('LOGO', style='LogoStyle').bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CNameStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(18)
    obj_font.name = 'Arial'

    p = document.add_paragraph()
    p.add_run('COMPANY NAME\n', style='CNameStyle').bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('DNameStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(18)
    obj_font.name = 'Arial'

    p = document.add_paragraph()
    p.add_run('\n\nDepartment Name\n\n\n', style='DNameStyle')
    p.add_run(f'\n\n{doc}\n\n\n', style='DNameStyle')
    # p.add_run('Document Name\n\n\n', style='DNameStyle')
    p.add_run(f'\n\nCreator: {doc.creator.user.first_name} {doc.creator.user.last_name}\n', style='DNameStyle')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_page_break()

    section = document.sections[0]
    section.different_first_page_header_footer = True
    header = section.header

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('TStyle', WD_STYLE_TYPE.TABLE)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Arial'

    table = header.add_table(2, 4, Inches(6))
    table.style = 'TStyle'

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('PStyle', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Arial'
    obj_font.bold = True

    row1_cells = table.rows[0].cells
    row1_cells[0].text = 'LOGO'
    row1_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row1_cells[0].paragraphs[0].runs[0].font.bold = True
    row1_cells[0].paragraphs[0].runs[0].font.size = Pt(20)
    row1_cells[1].text = 'Document Name:'
    row1_cells[1].paragraphs[0].runs[0].font.bold = True
    row1_cells[1].width = Inches(2.25)
    row1_cells[2].text = 'Version:'
    row1_cells[2].width = Inches(1.84)
    row1_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    row1_cells[3].text = str(doc.version)
    row1_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row1_cells[3].width = Inches(0.38)

    for cell in row1_cells:
        cell.vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_after = 0

    row2_cells = table.rows[1].cells
    row2_cells[1].text = str(doc)
    row2_cells[1].width = Inches(2.25)
    row2_cells[2].text = 'Running No:'
    row2_cells[2].width = Inches(1.84)
    row2_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    row2_cells[3].text = str(doc.running_no)
    row2_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row2_cells[3].width = Inches(0.38)

    row1_cells[0].merge(row2_cells[0])

    for cell in row2_cells:
        cell.vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_after = 0

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=download.docx'
    document.save(response)

    return response
